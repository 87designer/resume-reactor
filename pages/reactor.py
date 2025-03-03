import streamlit as st
import openai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import re
from typing import List, Optional


def generate_resume(job_description: str, resume_content: str, industry_key_word: str, api_key: str) -> str:
    """
    Calls OpenAI API to generate a tailored private-sector resume based on the job description and industry.

    Args:
        job_description (str): The job description provided by the user.
        resume_content (str): The user's federal resume content.
        industry_key_word (str): The industry type to tailor the resume.
        api_key (str): OpenAI API key for authentication.

    Returns:
        str: The formatted resume text generated by OpenAI.
    """
    client = openai.OpenAI(api_key=api_key)
    prompt = f"""
    I have a federal resume that I need to convert into a single-page private-sector resume tailored to a specific job. The goal is to maximize clarity, emphasize transferable skills, and ensure industry alignment while strictly adhering to space constraints.

    Inputs:
    1. Federal Resume: {resume_content}
    2. Job Description: {job_description}
    3. Target Industry: {industry_key_word}

    Strict Formatting & Length Constraints:
    - The entire resume must fit within one page in Arial 11pt font with standard 0.75-inch margins.
    - Use concise bullet points (ideally under two lines each) to maximize space.
    - Ensure the total word count does not exceed approximately 500-600 words to fit within one page.
    - Do not include unnecessary sections (e.g., federal job series, supervisors, GS levels).
    - Summarize only the most relevant experience, avoiding excessive detail.
    - Avoid redundancy and condense phrases without losing key meaning.

    Transformation Requirements:
    - Convert government-specific language into private-sector terminology.
    - Prioritize quantifiable achievements and key impact statements.
    - Maintain an ATS-friendly structure with industry-standard wording.
    - Remove all response context and output only the resume text with markdown formatting.

    Expected Output (Single Page, ATS-Friendly Format):
    - Professional Summary: A 2-3 sentence overview aligning with the job.
    - Key Skills: 6-8 bullet points summarizing relevant expertise.
    - Experience Section:
        - Limit to 2-3 most relevant jobs (or 10-15 years max).
        - Each role should have only 3-5 bullets, focusing on key achievements.
    - Education & Certifications: Keep brief, listing only essential details.

    Additional Instructions:
    - Do not include name, contact details, or headers that waste space.
    - Apply markdown formatting so headers and bullet points translate cleanly for export.
    - Ensure readability, prioritizing impactful, concise language over dense text.
    
    Output must be structured in a clean, professional format that fits entirely on a single page using Arial 11pt font.
    """
    
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an ATS optimization specialist skilled in converting federal resumes into concise, industry-ready formats."},
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content

def create_word_doc(resume_text: str, user_name: str, contact_info: List[str]) -> str:
    """
    Creates a formatted Word document from the generated resume text.

    Args:
        resume_text (str): The generated resume content.
        user_name (str): The user's full name.
        contact_info (List[str]): A list of the user's contact information.

    Returns:
        str: The file path to the generated Word document.
    """
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Add proper header
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    name_run = header_paragraph.add_run(user_name + "\n")
    name_run.bold = True
    name_run.font.size = Pt(24)
    
    contact_line = " | ".join(contact_info)
    header_paragraph.add_run(contact_line + "\n")
    
    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_path)
    return temp_path


def show_resume_reactor():
    st.title("🚀 ResumeReactor: Power Up Your Federal Resume")
    st.markdown(
        """
        Welcome to **ResumeReactor**, the tool that refines your federal resume to match job descriptions with precision.
        Simply upload your resume, paste a job description, and let AI generate a **polished, ATS-friendly** resume.
        """
    )

    # API Key Input
    openai_api_key: str = st.text_input("Enter your OpenAI API Key:", type="password")

    # User Information Inputs
    user_name: str = st.text_input("Enter Your Full Name:")
    email: str = st.text_input("Enter Your Email:")
    phone: str = st.text_input("Enter Your Phone Number:")
    linkedin: str = st.text_input("Enter Your LinkedIn Profile:")
    github: str = st.text_input("Enter Your GitHub Profile:")
    website: str = st.text_input("Enter Your Personal Website:")

    # Build contact info list dynamically, ignoring empty fields
    contact_info: List[str] = [info for info in [email, phone, linkedin, github, website] if info]

    # Resume Upload
    uploaded_file = st.file_uploader("📂 Drag and drop your federal resume (TXT, DOCX)", type=["txt", "docx"])

    # Job Description Input
    industry_key_word: str = st.text_input("Specify the industry of the job posting (e.g., IT, Finance, Healthcare, etc.)")
    job_description: str = st.text_area("📝 Paste the job description here:")

    # Persistent resume preview state
    if "generated_resume" not in st.session_state:
        st.session_state.generated_resume = ""

    # Resume Generation Button
    if st.button("Generate Resume"):
        if not user_name or not contact_info:
            st.error("Please fill in your full name and at least one contact detail.")
        else:
            if uploaded_file:
                if uploaded_file.type == "text/plain":
                    resume_content = uploaded_file.read().decode("utf-8")
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = Document(uploaded_file)
                    resume_content = "\n".join([para.text for para in doc.paragraphs])
                else:
                    st.error("Unsupported file format. Please upload TXT or DOCX.")
                    resume_content = ""
                
                if job_description and resume_content:
                    st.session_state.generated_resume = generate_resume(job_description, resume_content, industry_key_word, openai_api_key)

    # Display the resume preview persistently
    if st.session_state.generated_resume:
        st.subheader("🚀 AI-Powered Resume Output")
        st.markdown(st.session_state.generated_resume)
        
        file_path = create_word_doc(st.session_state.generated_resume, user_name, contact_info)
        with open(file_path, "rb") as file:
            st.download_button("💾 Download Word Document", data=file, file_name="formatted_resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
